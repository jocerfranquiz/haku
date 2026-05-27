#!/usr/bin/env python3
"""Generate a mixed test corpus for step 5 end-to-end indexing tests.

Creates healthy files across every supported extension, plus failure cases:
zero-byte PDF, encrypted PDF, mojibake docx, and a Spanish-accented doc.
Run once; output lands in engine/fixtures/corpus/.

Not part of the runtime — dev-only.
"""

from __future__ import annotations

from pathlib import Path

CORPUS_DIR = Path(__file__).resolve().parent / "fixtures" / "corpus"


def _write(name: str, content: str) -> None:
    (CORPUS_DIR / name).write_text(content, encoding="utf-8")


def _write_bytes(name: str, data: bytes) -> None:
    (CORPUS_DIR / name).write_bytes(data)


def gen_txt() -> None:
    _write(
        "readme.txt",
        "This is a plain text file used for testing haku indexing.\n"
        "It contains multiple lines of English prose.\n\n"
        "Search engines need to handle plain text gracefully.\n",
    )


def gen_md() -> None:
    _write(
        "notes.md",
        "# Research Notes\n\n"
        "## Introduction\n\n"
        "These are notes about information retrieval.\n"
        "Semantic search uses dense vector representations.\n\n"
        "## Methods\n\n"
        "We compare BM25 with dense retrieval approaches.\n"
        "Hybrid methods combine both lexical and semantic signals.\n",
    )


def gen_spanish_md() -> None:
    _write(
        "notas_es.md",
        "# Notas de Investigación\n\n"
        "## Introducción\n\n"
        "La búsqueda semántica utiliza representaciones vectoriales densas.\n"
        "Los motores de búsqueda modernos combinan señales léxicas y semánticas.\n\n"
        "## Métodos\n\n"
        "Comparamos BM25 con enfoques de recuperación densa.\n"
        "El niño estudió en el café de la universidad de Córdoba.\n",
    )


def gen_html() -> None:
    _write(
        "page.html",
        "<!DOCTYPE html>\n<html><head><title>Test Page</title></head>\n"
        "<body>\n<h1>Welcome</h1>\n"
        "<p>This is a test HTML page for indexing.</p>\n"
        "<p>It has multiple paragraphs with different content.</p>\n"
        "<h2>Section Two</h2>\n"
        "<p>More content here about search engines and retrieval.</p>\n"
        "</body></html>\n",
    )


def gen_docx() -> None:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test DOCX file for haku indexing.")
    doc.add_heading("Background", level=2)
    doc.add_paragraph(
        "Dense retrieval models encode queries and documents "
        "into high-dimensional vector spaces."
    )
    doc.save(str(CORPUS_DIR / "report.docx"))


def gen_pdf() -> None:
    import pymupdf  # type: ignore[import-untyped]

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "This is a test PDF for haku indexing.\n\n"
        "It contains text about information retrieval and search.",
    )
    doc.save(str(CORPUS_DIR / "article.pdf"))
    doc.close()


def gen_zero_byte_pdf() -> None:
    _write_bytes("empty.pdf", b"")


def gen_encrypted_pdf() -> None:
    import pymupdf  # type: ignore[import-untyped]

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This content is behind a password.")
    doc.save(
        str(CORPUS_DIR / "encrypted.pdf"),
        encryption=pymupdf.PDF_ENCRYPT_AES_256,
        user_pw="secret",
        owner_pw="owner",
    )
    doc.close()


def gen_mojibake_docx() -> None:
    """Create a .docx whose XML is structurally valid but contains mojibake:
    UTF-8 bytes of a Spanish string decoded as Latin-1, producing garbled text.
    mammoth should still parse it without crashing."""
    from docx import Document  # type: ignore[import-untyped]

    spanish = "El niño comió manzanas en el café"
    mojibake = spanish.encode("utf-8").decode("latin-1")
    doc = Document()
    doc.add_paragraph(mojibake)
    doc.save(str(CORPUS_DIR / "mojibake.docx"))


def gen_unsupported() -> None:
    _write("code.py", "# This .py file should be silently skipped.\nprint('hello')\n")
    _write("data.csv", "col1,col2\n1,2\n3,4\n")


def main() -> None:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    gen_txt()
    gen_md()
    gen_spanish_md()
    gen_html()
    gen_docx()
    gen_pdf()
    gen_zero_byte_pdf()
    gen_encrypted_pdf()
    gen_mojibake_docx()
    gen_unsupported()
    print(f"Corpus generated at {CORPUS_DIR}")
    for f in sorted(CORPUS_DIR.iterdir()):
        print(f"  {f.name} ({f.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
